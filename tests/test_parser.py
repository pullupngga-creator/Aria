"""Unit tests for document parser functions."""

from pathlib import Path

import pandas as pd
import pytest
from docx import Document as DocxDocument

from aria.document.parser import (
    extract_csv,
    extract_docx,
    extract_markdown,
    extract_pdf,
    extract_text,
    extract_txt,
    extract_xlsx,
    sanitize_text,
)
from aria.document.tokenizer import count_tokens, count_words
from aria.exceptions import DocumentParseError


class TestSanitizeText:
    """Tests for text sanitization."""

    def test_removes_null_bytes(self) -> None:
        """Test that null bytes are removed from text."""
        text = "Hello\x00World\x00Test"
        result = sanitize_text(text)
        assert "\x00" not in result
        assert result == "HelloWorldTest"

    def test_normalizes_line_endings(self) -> None:
        """Test that line endings are normalized to Unix style."""
        text = "Line1\r\nLine2\rLine3\n"
        result = sanitize_text(text)
        assert "\r\n" not in result
        assert "\r" not in result
        assert result == "Line1\nLine2\nLine3\n"

    def test_preserves_valid_text(self) -> None:
        """Test that valid text is preserved."""
        text = "Hello World\nThis is a test."
        result = sanitize_text(text)
        assert result == text


class TestExtractTxt:
    """Tests for TXT file extraction."""

    def test_extract_simple_txt(self, tmp_path: Path) -> None:
        """Test extracting text from a simple TXT file."""
        test_file = tmp_path / "test.txt"
        test_file.write_text("Hello World\nThis is a test.", encoding="utf-8")
        
        result = extract_txt(test_file)
        assert result == "Hello World\nThis is a test."

    def test_extract_txt_with_unicode(self, tmp_path: Path) -> None:
        """Test extracting text with Unicode characters."""
        test_file = tmp_path / "test.txt"
        test_file.write_text("Hello 世界 🌍", encoding="utf-8")
        
        result = extract_txt(test_file)
        assert "世界" in result
        assert "🌍" in result

    def test_extract_txt_invalid_encoding(self, tmp_path: Path) -> None:
        """Test that invalid encoding raises DocumentParseError."""
        test_file = tmp_path / "test.txt"
        # Write with latin-1 encoding but try to read as UTF-8
        test_file.write_bytes(b"\xff\xfe Invalid UTF-8")
        
        with pytest.raises(DocumentParseError):
            extract_txt(test_file)

    def test_extract_txt_sanitizes(self, tmp_path: Path) -> None:
        """Test that extracted text is sanitized."""
        test_file = tmp_path / "test.txt"
        test_file.write_text("Hello\r\nWorld\x00", encoding="utf-8")
        
        result = extract_txt(test_file)
        assert "\r\n" not in result
        assert "\x00" not in result


class TestExtractPdf:
    """Tests for PDF file extraction."""

    def test_extract_pdf_requires_valid_file(self, tmp_path: Path) -> None:
        """Test that invalid PDF raises DocumentParseError."""
        test_file = tmp_path / "invalid.pdf"
        test_file.write_bytes(b"Not a valid PDF")
        
        with pytest.raises(DocumentParseError):
            extract_pdf(test_file)

    def test_extract_pdf_missing_file(self, tmp_path: Path) -> None:
        """Test that missing file raises DocumentParseError."""
        test_file = tmp_path / "nonexistent.pdf"
        
        with pytest.raises(DocumentParseError):
            extract_pdf(test_file)


class TestCountWords:
    """Tests for word counting."""

    def test_count_simple_text(self) -> None:
        """Test counting words in simple text."""
        text = "Hello world this is a test"
        result = count_words(text)
        assert result == 6

    def test_count_words_with_newlines(self) -> None:
        """Test counting words with newlines."""
        text = "Hello\nworld\nthis\nis\na\ntest"
        result = count_words(text)
        assert result == 6

    def test_count_words_empty(self) -> None:
        """Test counting words in empty text."""
        text = ""
        result = count_words(text)
        assert result == 0

    def test_count_words_multiple_spaces(self) -> None:
        """Test that multiple spaces are handled correctly."""
        text = "Hello    world   this  is  test"
        result = count_words(text)
        assert result == 5


class TestCountTokens:
    """Tests for token counting."""

    def test_count_simple_text(self) -> None:
        """Test counting tokens in simple text."""
        text = "Hello world"
        result = count_tokens(text)
        assert result > 0
        assert result <= len(text)  # Tokens should be <= characters

    def test_count_tokens_empty(self) -> None:
        """Test counting tokens in empty text."""
        text = ""
        result = count_tokens(text)
        assert result == 0

    def test_count_tokens_longer_text(self) -> None:
        """Test counting tokens in longer text."""
        text = "This is a longer text to test token counting. It should have more tokens than characters divided by four."
        result = count_tokens(text)
        assert result > 0

    def test_count_tokens_unicode(self) -> None:
        """Test counting tokens with Unicode characters."""
        text = "Hello 世界 🌍"
        result = count_tokens(text)
        assert result > 0


# ── DOCX ─────────────────────────────────────────────────────────────────────


class TestExtractDocx:
    """Tests for DOCX file extraction."""

    def _make_docx(self, tmp_path: Path, paragraphs: list[str]) -> Path:
        """Helper: create a minimal DOCX with the given paragraphs."""
        doc = DocxDocument()
        for text in paragraphs:
            doc.add_paragraph(text)
        out = tmp_path / "test.docx"
        doc.save(str(out))
        return out

    def test_extracts_paragraphs(self, tmp_path: Path) -> None:
        """extract_docx returns body paragraph text."""
        path = self._make_docx(tmp_path, ["Hello World", "Second paragraph"])
        result = extract_docx(path)
        assert "Hello World" in result
        assert "Second paragraph" in result

    def test_skips_empty_paragraphs(self, tmp_path: Path) -> None:
        """Empty paragraphs are not included in output."""
        path = self._make_docx(tmp_path, ["Content", "", ""])
        result = extract_docx(path)
        # Should not contain triple newlines from blank paragraphs
        assert "Content" in result
        assert result.strip() != ""

    def test_extracts_table_cells(self, tmp_path: Path) -> None:
        """extract_docx includes table cell text separated by ' | '."""
        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Score"
        table.cell(1, 0).text = "Alice"
        table.cell(1, 1).text = "95"
        path = tmp_path / "table.docx"
        doc.save(str(path))

        result = extract_docx(path)
        assert "Name" in result
        assert "Score" in result
        assert "Alice" in result
        assert "95" in result

    def test_sanitizes_output(self, tmp_path: Path) -> None:
        """Output is sanitized (line endings normalised).

        Note: python-docx encodes to XML which forbids null bytes (\\x00),
        so null-byte sanitisation is verified at the sanitize_text unit level.
        Here we confirm \\r\\n is collapsed to \\n after extraction.
        """
        doc = DocxDocument()
        # Add a paragraph whose text contains a literal \r\n sequence
        # (the run text is set directly to bypass paragraph-level newline handling)
        para = doc.add_paragraph()
        run = para.add_run("Line\r\nBreak")
        path = tmp_path / "dirty.docx"
        doc.save(str(path))

        result = extract_docx(path)
        assert "\r\n" not in result

    def test_raises_on_invalid_file(self, tmp_path: Path) -> None:
        """extract_docx raises DocumentParseError for a non-DOCX file."""
        bad = tmp_path / "bad.docx"
        bad.write_bytes(b"Not a DOCX file")
        with pytest.raises(DocumentParseError):
            extract_docx(bad)


# ── CSV ──────────────────────────────────────────────────────────────────────


class TestExtractCsv:
    """Tests for CSV file extraction."""

    def test_extracts_headers_and_rows(self, tmp_path: Path) -> None:
        """extract_csv returns column names and row data."""
        csv_file = tmp_path / "data.csv"
        csv_file.write_text("name,age\nAlice,30\nBob,25", encoding="utf-8")
        result = extract_csv(csv_file)
        assert "name" in result
        assert "Alice" in result
        assert "Bob" in result

    def test_latin1_fallback(self, tmp_path: Path) -> None:
        """extract_csv falls back to latin-1 when UTF-8 decoding fails."""
        csv_file = tmp_path / "latin.csv"
        # Write latin-1 encoded data (ñ is 0xF1 in latin-1, invalid UTF-8 start)
        csv_file.write_bytes("col\nca\xf1on".encode("latin-1"))
        result = extract_csv(csv_file)
        assert "col" in result

    def test_numeric_data_preserved(self, tmp_path: Path) -> None:
        """Numeric columns are included as-is without rounding."""
        csv_file = tmp_path / "nums.csv"
        csv_file.write_text("x,y\n1.234,5.678\n9.0,0.1", encoding="utf-8")
        result = extract_csv(csv_file)
        assert "1.234" in result
        assert "5.678" in result

    def test_sanitizes_output(self, tmp_path: Path) -> None:
        """Output is sanitized."""
        csv_file = tmp_path / "dirty.csv"
        csv_file.write_text("col\nval\x00ue", encoding="utf-8")
        result = extract_csv(csv_file)
        assert "\x00" not in result

    def test_raises_on_invalid_file(self, tmp_path: Path) -> None:
        """extract_csv raises DocumentParseError for an unreadable file."""
        missing = tmp_path / "missing.csv"
        with pytest.raises(DocumentParseError):
            extract_csv(missing)


# ── XLSX ─────────────────────────────────────────────────────────────────────


class TestExtractXlsx:
    """Tests for Excel XLSX file extraction."""

    def _make_xlsx(self, tmp_path: Path, sheets: dict[str, pd.DataFrame]) -> Path:
        """Helper: write a multi-sheet XLSX file."""
        path = tmp_path / "test.xlsx"
        with pd.ExcelWriter(str(path), engine="openpyxl") as writer:
            for sheet_name, df in sheets.items():
                df.to_excel(writer, sheet_name=sheet_name, index=False)
        return path

    def test_extracts_single_sheet(self, tmp_path: Path) -> None:
        """extract_xlsx returns data from a single-sheet workbook."""
        df = pd.DataFrame({"Name": ["Alice", "Bob"], "Score": [95, 88]})
        path = self._make_xlsx(tmp_path, {"Results": df})
        result = extract_xlsx(path)
        assert "Name" in result
        assert "Alice" in result
        assert "95" in result

    def test_extracts_multiple_sheets(self, tmp_path: Path) -> None:
        """extract_xlsx concatenates all sheets with a ## Sheet: heading."""
        df1 = pd.DataFrame({"A": [1]})
        df2 = pd.DataFrame({"B": [2]})
        path = self._make_xlsx(tmp_path, {"First": df1, "Second": df2})
        result = extract_xlsx(path)
        assert "## Sheet: First" in result
        assert "## Sheet: Second" in result

    def test_skips_empty_sheets(self, tmp_path: Path) -> None:
        """Empty sheets are silently skipped."""
        df_real = pd.DataFrame({"Col": ["value"]})
        df_empty = pd.DataFrame()
        path = self._make_xlsx(tmp_path, {"Data": df_real, "Empty": df_empty})
        result = extract_xlsx(path)
        assert "## Sheet: Data" in result
        assert "## Sheet: Empty" not in result

    def test_raises_on_invalid_file(self, tmp_path: Path) -> None:
        """extract_xlsx raises DocumentParseError for a non-XLSX file."""
        bad = tmp_path / "bad.xlsx"
        bad.write_bytes(b"Not an Excel file")
        with pytest.raises(DocumentParseError):
            extract_xlsx(bad)


# ── Markdown ──────────────────────────────────────────────────────────────────


class TestExtractMarkdown:
    """Tests for Markdown file extraction."""

    def test_extracts_raw_source(self, tmp_path: Path) -> None:
        """extract_markdown returns the raw Markdown source unchanged."""
        content = "# Title\n\nSome **bold** text and a [link](http://example.com)."
        md_file = tmp_path / "doc.md"
        md_file.write_text(content, encoding="utf-8")
        result = extract_markdown(md_file)
        assert "# Title" in result
        assert "**bold**" in result
        assert "[link](http://example.com)" in result

    def test_unicode_content(self, tmp_path: Path) -> None:
        """extract_markdown handles Unicode characters correctly."""
        md_file = tmp_path / "unicode.md"
        md_file.write_text("# 日本語\n\nContent 🌍", encoding="utf-8")
        result = extract_markdown(md_file)
        assert "日本語" in result
        assert "🌍" in result

    def test_sanitizes_output(self, tmp_path: Path) -> None:
        """Output is sanitized (null bytes, line endings)."""
        md_file = tmp_path / "dirty.md"
        md_file.write_text("Hello\r\nWorld\x00", encoding="utf-8")
        result = extract_markdown(md_file)
        assert "\x00" not in result
        assert "\r\n" not in result

    def test_raises_on_missing_file(self, tmp_path: Path) -> None:
        """extract_markdown raises DocumentParseError when file is missing."""
        with pytest.raises(DocumentParseError):
            extract_markdown(tmp_path / "missing.md")


# ── Dispatcher ────────────────────────────────────────────────────────────────


class TestExtractTextDispatcher:
    """Tests for the unified extract_text dispatcher."""

    def test_routes_txt(self, tmp_path: Path) -> None:
        """Dispatcher routes .txt to extract_txt."""
        f = tmp_path / "file.txt"
        f.write_text("plain text", encoding="utf-8")
        assert extract_text(f) == "plain text"

    def test_routes_markdown(self, tmp_path: Path) -> None:
        """Dispatcher routes .md to extract_markdown."""
        f = tmp_path / "readme.md"
        f.write_text("# Heading", encoding="utf-8")
        result = extract_text(f)
        assert "# Heading" in result

    def test_routes_csv(self, tmp_path: Path) -> None:
        """Dispatcher routes .csv to extract_csv."""
        f = tmp_path / "data.csv"
        f.write_text("a,b\n1,2", encoding="utf-8")
        result = extract_text(f)
        assert "a" in result
        assert "1" in result

    def test_routes_xlsx(self, tmp_path: Path) -> None:
        """Dispatcher routes .xlsx to extract_xlsx."""
        df = pd.DataFrame({"X": [42]})
        path = tmp_path / "data.xlsx"
        df.to_excel(str(path), index=False)
        result = extract_text(path)
        assert "42" in result

    def test_routes_docx(self, tmp_path: Path) -> None:
        """Dispatcher routes .docx to extract_docx."""
        doc = DocxDocument()
        doc.add_paragraph("Dispatcher test")
        path = tmp_path / "file.docx"
        doc.save(str(path))
        result = extract_text(path)
        assert "Dispatcher test" in result

    def test_raises_on_unsupported_extension(self, tmp_path: Path) -> None:
        """Dispatcher raises DocumentParseError for unknown extensions."""
        f = tmp_path / "file.pptx"
        f.write_bytes(b"not supported")
        with pytest.raises(DocumentParseError, match="Unsupported file type"):
            extract_text(f)

    def test_case_insensitive_extension(self, tmp_path: Path) -> None:
        """Dispatcher handles upper-case extensions (e.g. .TXT, .MD)."""
        f = tmp_path / "readme.MD"
        f.write_text("# Upper MD", encoding="utf-8")
        result = extract_text(f)
        assert "# Upper MD" in result
