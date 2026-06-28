"""Document text extraction parsers for PDF, TXT, DOCX, CSV, XLSX, and Markdown."""

import logging
from pathlib import Path

import fitz  # type: ignore # pymupdf
import pandas as pd
from docx import Document as DocxDocument

from aria.exceptions import DocumentParseError

logger = logging.getLogger(__name__)

# Supported extensions mapped to their extractor (for the dispatcher)
_SUPPORTED_EXTENSIONS: frozenset[str] = frozenset(
    {".pdf", ".txt", ".docx", ".csv", ".xlsx", ".md"}
)


def sanitize_text(text: str) -> str:
    """Sanitize extracted text by removing null bytes and normalizing line endings.

    Args:
        text: Raw extracted text.

    Returns:
        Sanitized text with Unix line endings and no null bytes.
    """
    text = text.replace("\x00", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return text


def extract_pdf(file_path: Path) -> str:
    """Extract text from PDF file using pymupdf.

    Args:
        file_path: Path to the PDF file.

    Returns:
        Extracted text as a string.

    Raises:
        DocumentParseError: If PDF extraction fails.
    """
    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return sanitize_text(text)
    except fitz.FileDataError as e:
        logger.error("PDF file is corrupted or invalid", extra={"path": str(file_path)})
        raise DocumentParseError(f"PDF file is corrupted: {file_path}") from e
    except Exception as e:
        logger.error("Failed to extract text from PDF", extra={"path": str(file_path)})
        raise DocumentParseError(f"Failed to extract text from PDF: {file_path}") from e


def extract_txt(file_path: Path) -> str:
    """Extract text from plain text file.

    Args:
        file_path: Path to the TXT file.

    Returns:
        Extracted text as a string.

    Raises:
        DocumentParseError: If text extraction fails.
    """
    try:
        with open(file_path, encoding="utf-8") as f:
            text = f.read()
        return sanitize_text(text)
    except UnicodeDecodeError as e:
        logger.error("Failed to decode text file as UTF-8", extra={"path": str(file_path)})
        raise DocumentParseError(f"Text file encoding error: {file_path}") from e
    except Exception as e:
        logger.error("Failed to read text file", extra={"path": str(file_path)})
        raise DocumentParseError(f"Failed to read text file: {file_path}") from e


def extract_docx(file_path: Path) -> str:
    """Extract text from a DOCX file using python-docx.

    Captures text from both body paragraphs and table cells so that
    structured content (e.g. Word tables) is included in the extraction.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Extracted text joined by double newlines.

    Raises:
        DocumentParseError: If DOCX extraction fails.
    """
    try:
        doc = DocxDocument(str(file_path))
        parts: list[str] = []

        # Body paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Table cells (row-by-row, cell-by-cell)
        for table in doc.tables:
            for row in table.rows:
                row_parts: list[str] = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_parts.append(cell_text)
                if row_parts:
                    parts.append(" | ".join(row_parts))

        extracted = "\n\n".join(parts)
        logger.info(
            "DOCX extracted successfully",
            extra={"path": str(file_path), "length": len(extracted)},
        )
        return sanitize_text(extracted)
    except Exception as e:
        logger.error("Failed to extract text from DOCX", extra={"path": str(file_path)})
        raise DocumentParseError(f"Failed to extract text from DOCX: {file_path}") from e


def extract_csv(file_path: Path) -> str:
    """Extract text from a CSV file using pandas.

    Tries UTF-8 encoding first; falls back to latin-1 for files exported
    from legacy tools (e.g. Excel on Windows).

    Args:
        file_path: Path to the CSV file.

    Returns:
        String representation of the data frame with headers.

    Raises:
        DocumentParseError: If CSV extraction fails.
    """
    try:
        try:
            df = pd.read_csv(file_path, encoding="utf-8")
        except UnicodeDecodeError:
            logger.warning(
                "CSV UTF-8 decode failed, falling back to latin-1",
                extra={"path": str(file_path)},
            )
            df = pd.read_csv(file_path, encoding="latin-1")

        extracted = df.to_string(index=False)
        logger.info(
            "CSV extracted successfully",
            extra={"path": str(file_path), "rows": len(df), "cols": len(df.columns)},
        )
        return sanitize_text(extracted)
    except DocumentParseError:
        raise
    except Exception as e:
        logger.error("Failed to extract text from CSV", extra={"path": str(file_path)})
        raise DocumentParseError(f"Failed to extract text from CSV: {file_path}") from e


def extract_xlsx(file_path: Path) -> str:
    """Extract text from an Excel XLSX file using pandas.

    All sheets are extracted and concatenated with a ``## Sheet: <name>``
    heading prefix. Empty sheets are skipped with a debug log.

    Args:
        file_path: Path to the XLSX file.

    Returns:
        Concatenated string of all non-empty sheet contents.

    Raises:
        DocumentParseError: If XLSX extraction fails.
    """
    try:
        xl = pd.ExcelFile(file_path)
        parts: list[str] = []

        for sheet_name in xl.sheet_names:
            df = xl.parse(sheet_name)
            if df.empty:
                logger.debug(
                    "Skipping empty sheet",
                    extra={"sheet": sheet_name, "path": str(file_path)},
                )
                continue
            sheet_text = df.to_string(index=False)
            parts.append(f"## Sheet: {sheet_name}\n\n{sheet_text}")

        if not parts:
            logger.warning("XLSX file has no non-empty sheets", extra={"path": str(file_path)})
            return ""

        extracted = "\n\n".join(parts)
        logger.info(
            "XLSX extracted successfully",
            extra={"path": str(file_path), "sheets": len(parts)},
        )
        return sanitize_text(extracted)
    except Exception as e:
        logger.error("Failed to extract text from XLSX", extra={"path": str(file_path)})
        raise DocumentParseError(f"Failed to extract text from XLSX: {file_path}") from e


def extract_markdown(file_path: Path) -> str:
    """Extract text from a Markdown file.

    Markdown source is injected into the LLM context as-is — the model
    understands Markdown syntax natively, so no conversion is needed.

    Args:
        file_path: Path to the Markdown (.md) file.

    Returns:
        Raw Markdown source text.

    Raises:
        DocumentParseError: If the file cannot be read.
    """
    try:
        with open(file_path, encoding="utf-8") as f:
            text = f.read()
        logger.info(
            "Markdown extracted successfully",
            extra={"path": str(file_path), "length": len(text)},
        )
        return sanitize_text(text)
    except UnicodeDecodeError as e:
        logger.error(
            "Failed to decode Markdown file as UTF-8", extra={"path": str(file_path)}
        )
        raise DocumentParseError(f"Markdown file encoding error: {file_path}") from e
    except Exception as e:
        logger.error("Failed to read Markdown file", extra={"path": str(file_path)})
        raise DocumentParseError(f"Failed to read Markdown file: {file_path}") from e


def extract_text(file_path: Path) -> str:
    """Unified dispatcher — route to the correct extractor based on file extension.

    This is the single entry point used by :class:`~aria.document.vault.VaultManager`
    and any other caller. Adding a new format only requires adding a new ``extract_*``
    function and mapping it here.

    Args:
        file_path: Path to the document.

    Returns:
        Extracted plain text ready for tokenisation and context injection.

    Raises:
        DocumentParseError: If the file type is unsupported or extraction fails.
    """
    ext = file_path.suffix.lower()
    dispatch: dict[str, object] = {
        ".pdf": extract_pdf,
        ".txt": extract_txt,
        ".docx": extract_docx,
        ".csv": extract_csv,
        ".xlsx": extract_xlsx,
        ".md": extract_markdown,
    }
    extractor = dispatch.get(ext)
    if extractor is None:
        raise DocumentParseError(
            f"Unsupported file type '{ext}'. "
            f"Supported: {', '.join(sorted(dispatch.keys()))}"
        )
    return extractor(file_path)  # type: ignore[operator]
